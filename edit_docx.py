import docx
import os

def main():
    try:
        source_path = r"C:\Users\quyen\Downloads\baocaoDACNCNPM1_Nhóm 6.docx"
        target_path = r"C:\xampp\htdocs\The Stationery Muse\Bao_Cao_Chinh_Thuc_The_Stationery_Muse.docx"
        
        # Define replacements
        replacements = {
            "QUẢN LÝ BÁN CÂY CẢNH": "BÁN VĂN PHÒNG PHẨM THE STATIONERY MUSE",
            "CÂY CẢNH": "VĂN PHÒNG PHẨM",
            "cây cảnh": "văn phòng phẩm",
            "Cây cảnh": "Văn phòng phẩm",
            "Cửa hàng cây cảnh": "Cửa hàng văn phòng phẩm",
            "cửa hàng cây cảnh": "cửa hàng văn phòng phẩm",
            "thú cưng": "sản phẩm",
            "Thú cưng": "Sản phẩm",
            "MEOMONK": "THE STATIONERY MUSE",
            "Productss": "Products",
            "Product_id": "id",
            "User_id": "id",
            "Category_id": "id",
            "Order_id": "id",
            "Contact_id": "id",
            "Post_id": "id",
            "Cây Xanh": "Sổ Tay Planner",
            "Cây Hoa Đậu Biếc": "Bút Bi Mực Nước",
            "Hạt Giống": "Tập Vở",
            "Chậu Cây": "Bút Dạ Quang",
            "Cây Bonsai": "Dụng cụ học tập"
        }

        # Open the original word file
        doc = docx.Document(source_path)
        
        for p in doc.paragraphs:
            text = p.text
            changed = False
            for old_t in replacements.keys():
                if old_t in text:
                    changed = True
                    break
            
            if changed:
                replaced_in_runs = False
                for run in p.runs:
                    for old_t, new_t in replacements.items():
                        if old_t in run.text:
                            run.text = run.text.replace(old_t, new_t)
                            replaced_in_runs = True
                
                if not replaced_in_runs:
                    new_text = p.text
                    for old_t, new_t in replacements.items():
                        new_text = new_text.replace(old_t, new_t)
                    if new_text != p.text:
                        for i in range(len(p.runs)):
                            p.runs[i].text = ""
                        p.add_run(new_text)

        # Do the same for tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text
                        changed = False
                        for old_t in replacements.keys():
                            if old_t in text:
                                changed = True
                                break
                        
                        if changed:
                            replaced_in_runs = False
                            for run in p.runs:
                                for old_t, new_t in replacements.items():
                                    if old_t in run.text:
                                        run.text = run.text.replace(old_t, new_t)
                                        replaced_in_runs = True
                            
                            if not replaced_in_runs:
                                new_text = p.text
                                for old_t, new_t in replacements.items():
                                    new_text = new_text.replace(old_t, new_t)
                                if new_text != p.text:
                                    for i in range(len(p.runs)):
                                        p.runs[i].text = ""
                                    p.add_run(new_text)
                                    
        doc.save(target_path)
        print(f"SUCCESS: Saved to {target_path}")
    except Exception as e:
        print(f"FAILED: {str(e)}")

if __name__ == '__main__':
    main()
