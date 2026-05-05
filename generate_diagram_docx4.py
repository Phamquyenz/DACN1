import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import os
import time

try:
    from plantuml import PlantUML
except Exception as e:
    print(e)

def download_diagram_plantuml(puml_text, puml_file, filename):
    print("Generating", filename)
    pl = PlantUML(url='http://www.plantuml.com/plantuml/png/')
    with open(puml_file, 'w', encoding='utf-8') as f:
        f.write(puml_text)
    try:
        pl.processes_file(puml_file, outfile=filename)
        time.sleep(2)  # Give the server a breather
    except Exception as e:
        print("Failed download:", e)
        pass

def create_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

def main():
    doc = docx.Document()
    doc.add_paragraph('\n\n\n')
    p = doc.add_paragraph('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ\nKHOA CÔNG NGHỆ THÔNG TIN')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    p_title1 = doc.add_paragraph('BÀI TẬP LỚN\nHỌC PHẦN: ĐỒ ÁN CHUYÊN NGÀNH CÔNG NGHỆ PHẦN MỀM')
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    p_title2 = doc.add_paragraph('TÊN ĐỀ TÀI: XÂY DỰNG WEBSITE QUẢN LÝ BÁN VĂN PHÒNG PHẨM "THE STATIONERY MUSE"')
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_title2.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph('\n\n\n\n\n\n\n\n\n')
    p_title3 = doc.add_paragraph('Sinh viên thực hiện: ...\nGiảng viên hướng dẫn: ...')
    p_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # MỤC LỤC
    doc.add_heading('MỤC LỤC', 1)
    toc = [
        "CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI",
        "  1.1. Lý do chọn đề tài",
        "  1.2. Mục tiêu của đề tài",
        "CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG (LARAVEL)",
        "CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ",
        "  3.1. Sơ đồ Usecase",
        "  3.2. Sơ đồ Hoạt động",
        "  3.3. Sơ đồ Tuần tự",
        "CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU & ERD",
        "CHƯƠNG V: KẾT LUẬN"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading('CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI', 1)
    doc.add_heading('1.1. Lý do chọn đề tài', 2)
    doc.add_paragraph('Trong thời đại số hóa, thương mại điện tử giúp doanh nghiệp kinh doanh văn phòng phẩm quản lý tồn kho (Stock), đơn hàng, và khuyến mãi hiệu quả hơn so với truyền thống. "The Stationery Muse" được xây dựng để cung cấp công cụ tự động hóa toàn diện từ mua sắm online đến các chương trình chiết khấu ưu đãi khách hàng.')
    doc.add_heading('1.2. Mục tiêu', 2)
    doc.add_paragraph('Xây dựng hệ thống bằng Laravel MVC tinh gọn, kết hợp giao diện với Tailwind CSS, tích hợp thanh toán (VNPAY), kiểm soát chính xác cơ chế áp dụng Mã giảm giá (Vouchers).')

    doc.add_heading('CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG', 1)
    doc.add_paragraph('Dự án dựa trên Framework Laravel 10 (được viết bằng PHP). Khung kiến trúc MVC đảm bảo duy trì logic phần mềm (Controllers) và các mô hình CSDL MySQL rõ ràng. Giao diện tùy ưu nhanh chóng chuyên nghiệp thông qua Tailwind CSS.')

    doc.add_page_break()
    doc.add_heading('CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ', 1)
    
    # Sơ đồ Usecase
    puml_usecase = """@startuml
left to right direction
actor "Khách hàng" as KH
actor "Admin" as AD
package "The Stationery Muse" {
  usecase "Đăng nhập" as UC1
  usecase "Giỏ hàng" as UC2
  usecase "Thanh toán (VNPAY)" as UC3
  usecase "Quản lý / Phân quyền" as UC8
  usecase "Quản lý Sản phẩm / Kho" as UC9
  usecase "Quản lý Vouchers" as UC10
  usecase "Xử lý Đơn hàng" as UC11
}
KH --> UC1
KH --> UC2
KH --> UC3
AD --> UC1
AD --> UC8
AD --> UC9
AD --> UC10
AD --> UC11
@enduml"""
    download_diagram_plantuml(puml_usecase, 'uc_temp.puml', "uc_diag.png")
    doc.add_heading('3.1 Sơ đồ Usecase Tổng quát', 2)
    if os.path.exists("uc_diag.png"):
        doc.add_picture("uc_diag.png", width=Inches(6.0))
        p = doc.add_paragraph('Hình 1: Biểu đồ Usecase hệ thống bán văn phòng phẩm')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Lỗi tải Sơ đồ Use case từ PlantUML]")

    # Sơ đồ Activity
    puml_activity = """@startuml
start
:Thêm Giỏ hàng & Checkout;
:Áp dụng Mã giảm giá (Voucher);
if (Voucher lệ?) then (Có)
  :Trừ theo Giá trị Voucher;
else (Không)
  :Thông báo Lỗi;
endif
:Trừ số lượng Tồn kho của Sản phẩm;
:Tạo Đơn Hàng vào MySQL;
:Mở cổng thanh toán VNPAY;
if (Thanh toán VNPAY?) then (Thành công)
  :Cập nhật trạng thái "Processing";
else (Huỷ / Thất bại)
  :Giao dịch huỷ, hoàn tồn kho;
endif
stop
@enduml"""
    download_diagram_plantuml(puml_activity, 'act_temp.puml', "act_diag.png")
    doc.add_heading('3.2 Sơ đồ Hoạt động (Quy trình Checkout & VNPAY)', 2)
    if os.path.exists("act_diag.png"):
        doc.add_picture("act_diag.png", width=Inches(4.5))
        p = doc.add_paragraph('Hình 2: Sơ đồ hoạt động xử lý vòng lặp đơn hàng')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sơ đồ Sequence
    puml_sequence = """@startuml
actor KH as "Khách hàng"
boundary UI as "Blade View"
control CC as "OrderController"
entity DB as "MySQL"

KH -> UI: Chọn thanh toán & Voucher
UI -> CC: POST Đơn hàng
CC -> DB: Truy vấn chi tiết Voucher
DB --> CC: Thông số giới hạn/hạn mức
alt Áp dụng thành công
    CC -> DB: Insert tbl_Orders & tbl_OrderItems
    DB --> CC: order_id
    CC -> DB: Update số lượng Sản phẩm (Stock)
    CC --> UI: Khởi tạo Redirect URL (VNPAY)
    UI --> KH: Cổng thanh toán VNPAY
else Hết lượt / Hết hạn
    CC --> UI: Exception / Flash Error
    UI --> KH: Trả thông báo Lỗi trên View
end
@enduml"""
    download_diagram_plantuml(puml_sequence, 'seq_temp.puml', "seq_diag.png")
    doc.add_heading('3.3 Sơ đồ Tuần tự (Xác thực Voucher)', 2)
    if os.path.exists("seq_diag.png"):
        doc.add_picture("seq_diag.png", width=Inches(6.0))
        p = doc.add_paragraph('Hình 3: Sơ đồ tuần tự tương tác Khách hàng với Voucher')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ERD 
    doc.add_heading('CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU & ERD', 1)
    puml_erd = """@startuml
entity "users" as U {
  * id : bigint
  --
  name : varchar
  role : varchar
}
entity "products" as P {
  * id : bigint
  --
  category_id : bigint
  name : varchar
  stock : int
  sale_price : decimal
  flash_sale_end : timestamp
}
entity "vouchers" as V {
  * id : bigint
  --
  code : varchar
  discount_amount : decimal
}
entity "orders" as O {
  * id : bigint
  --
  user_id : bigint
  voucher_id : bigint
  total_amount : decimal
  status : varchar
}
entity "order_items" as OI {
  * id : bigint
  --
  order_id : bigint
  product_id : bigint
  quantity : int
  price : decimal
}
U "1" -- "0..*" O : creates >
P "1" -- "0..*" OI : includes >
O "1" -- "1..*" OI : contains >
V "1" -- "0..*" O : affects >
@enduml"""
    download_diagram_plantuml(puml_erd, 'erd_temp.puml', "erd_diag.png")
    doc.add_heading('4.1 Sơ đồ Thực thể Liên kết (ERD)', 2)
    if os.path.exists("erd_diag.png"):
        doc.add_picture("erd_diag.png", width=Inches(5.0))
        p = doc.add_paragraph('Hình 4: Sơ đồ ERD Hệ thống The Stationery Muse')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.2 Thiết kế Migrations', 2)
    doc.add_paragraph('Bảng 1: Bảng products (' + 'Sản phẩm'.encode('utf8').decode('utf8') + ')')
    create_table(doc, ['Trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt', 'Khóa chính'],
        ['name', 'Varchar', 'Tên sản phẩm'],
        ['sale_price', 'Decimal', 'Khuyến mãi'],
        ['stock', 'Int', 'Lượng tồn kho thực tế']
    ])

    doc.add_paragraph('\nBảng 2: Bảng vouchers (' + 'Mã giảm giá'.encode('utf8').decode('utf8') + ')')
    create_table(doc, ['Trường', 'Kiểu', 'Ý nghĩa'], [
        ['code', 'Varchar', 'Mã hiển thị (Ví dụ: SALE50)'],
        ['discount_amount', 'Decimal', 'Mức giảm']
    ])

    doc.add_paragraph('\nBảng 3: Bảng order_items (' + 'Chi tiết'.encode('utf8').decode('utf8') + ')')
    create_table(doc, ['Trường', 'Kiểu', 'Ý nghĩa'], [
        ['order_id', 'BigInt', 'Liên kết hệ thống Đơn hàng'],
        ['product_id', 'BigInt', 'Liên kết hệ thống Sản phẩm'],
        ['quantity', 'Int', 'Số lượng xuất kho']
    ])

    doc.add_page_break()
    doc.add_heading('CHƯƠNG V: KẾT LUẬN', 1)
    doc.add_paragraph('Hệ thống đạt mục tiêu thiết lập Website bán hàng The Stationery Muse dựa trên kỹ thuật lập trình MVC bằng Framework Laravel hiện đại. Hệ thống giải quyết tốt bài toán trừ kho sản phẩm và ứng dụng mã giảm giá (Vouchers) tự động.')

    # Ghi đè file chính
    doc.save(r'C:\xampp\htdocs\The Stationery Muse\Báo_Cáo_Siêu_Hoàn_Chỉnh_Có_Hình.docx')

    # Dọn tập tin cũ
    for suffix in ["uc", "act", "seq", "erd"]:
        if os.path.exists(f"{suffix}_temp.puml"): os.remove(f"{suffix}_temp.puml")
        if os.path.exists(f"{suffix}_diag.png"): os.remove(f"{suffix}_diag.png")

if __name__ == '__main__':
    main()
