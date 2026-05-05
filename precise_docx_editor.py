import docx

def main():
    source_path = r"C:\Users\quyen\Downloads\baocaoDACNCNPM1_Nhóm 6.docx"
    target_path = r"C:\xampp\htdocs\The Stationery Muse\Bao_Cao_Hoan_Chinh_The_Stationery_Muse.docx"
    
    try:
        doc = docx.Document(source_path)
    except Exception as e:
        print("Error opening doc:", e)
        return

    # 1. Global replacements
    replacements = {
        "QUẢN LÝ BÁN CÂY CẢNH": "QUẢN LÝ BÁN VĂN PHÒNG PHẨM",
        "CÂY CẢNH": "VĂN PHÒNG PHẨM",
        "cây cảnh": "văn phòng phẩm",
        "Cây cảnh": "Văn phòng phẩm",
        "Cửa hàng cây cảnh": "Cửa hàng văn phòng phẩm",
        "cửa hàng cây cảnh": "cửa hàng văn phòng phẩm",
        "thú cưng": "sản phẩm",
        "Thú cưng": "Sản phẩm",
        "PHP và MySQL": "Framework Laravel và MySQL",
        "thuần PHP": "Laravel",
        "HTML/CSS/SCSS và JavaScript": "Tailwind CSS, Alpine.js và Laravel Blade",
        "MEOMONK": "THE STATIONERY MUSE",
        '"Xanh Lá"': '"The Stationery Muse"',
        "Cây Xanh": "Sổ Tay Planner",
        "Cây Hoa Đậu Biếc": "Bút Bi Mực Nước",
        "Productss": "Products (Sản phẩm)",
        "Contacts": "Order_Items (Chi tiết Đơn hàng)",
        "Blog_posts": "Vouchers (Mã giảm giá)",
        "Blog": "Voucher",
        "diễn đàn": "hệ thống",
        "Diễn đàn": "Hệ thống"
    }

    # Replace in paragraphs
    for p in doc.paragraphs:
        changed = False
        for old_t in replacements.keys():
            if old_t in p.text:
                changed = True
                break
        if changed:
            # simple text replace across the whole paragraph, clearing runs
            new_text = p.text
            for old_t, new_t in replacements.items():
                new_text = new_text.replace(old_t, new_t)
            # preserve basic alignment/style
            style = p.style
            for i in range(len(p.runs)):
                p.runs[i].text = ""
            p.add_run(new_text)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    changed = False
                    for old_t in replacements.keys():
                        if old_t in p.text:
                            changed = True
                            break
                    if changed:
                        new_text = p.text
                        for old_t, new_t in replacements.items():
                            new_text = new_text.replace(old_t, new_t)
                        for i in range(len(p.runs)):
                            p.runs[i].text = ""
                        p.add_run(new_text)

    # 2. Hardcode table modifications for DB schema 
    # Table 4 is likely Blog_posts -> we change to Vouchers
    # We will search for tables by their headers or first row content
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) >= 4:
            first_row_text = "".join([c.text for c in table.rows[1].cells])
            
            # Map Vouchers (was Blog_posts)
            if "Post_id" in first_row_text or "Mã bài viết" in first_row_text:
                # Rewrite to Voucher table
                table_data = [
                    ["1","id","BigInt","PRIMARY KEY","Mã Voucher"],
                    ["2","code","Varchar(50)","NOT NULL","Mã giảm giá"],
                    ["3","discount_amount","Decimal(10,2)","NOT NULL","Giá trị giảm"],
                    ["4","usage_limit","Int","NULLABLE","Giới hạn sử dụng"],
                    ["5","expires_at","Timestamp","NULLABLE","Hạn sử dụng"]
                ]
                # clear existing rows except header
                while len(table.rows) > 1:
                    # remove row
                    tbl = table._tbl
                    tr = table.rows[1]._tr
                    tbl.remove(tr)
                # add new rows
                for r_data in table_data:
                    row = table.add_row()
                    for i, val in enumerate(r_data):
                        if i < len(row.cells):
                            row.cells[i].text = val

            # Map Order_Items (was Contacts)
            elif "Contact_id" in first_row_text or "Mã liên hệ" in first_row_text:
                table_data = [
                    ["1","id","BigInt","PRIMARY KEY","Mã chi tiết"],
                    ["2","order_id","BigInt","FOREIGN KEY","Mã đơn hàng"],
                    ["3","product_id","BigInt","FOREIGN KEY","Mã sản phẩm"],
                    ["4","quantity","Int","NOT NULL","Số lượng"],
                    ["5","price","Decimal(10,2)","NOT NULL","Giá bán"]
                ]
                while len(table.rows) > 1:
                    tbl = table._tbl
                    tr = table.rows[1]._tr
                    tbl.remove(tr)
                for r_data in table_data:
                    row = table.add_row()
                    for i, val in enumerate(r_data):
                        if i < len(row.cells):
                            row.cells[i].text = val
            
            # Map Products (was Productss)
            elif "Product_id" in first_row_text or "Mã sản phẩm" in first_row_text:
                table_data = [
                    ["1","id","BigInt","PRIMARY KEY","Mã sản phẩm"],
                    ["2","category_id","BigInt","FOREIGN KEY","Mã danh mục"],
                    ["3","name","Varchar(255)","NOT NULL","Tên sản phẩm"],
                    ["4","price","Decimal(10,2)","NOT NULL","Giá bán"],
                    ["5","sale_price","Decimal(10,2)","NULLABLE","Giá khuyến mãi"],
                    ["6","stock","Int","NOT NULL","Số lượng kho"],
                    ["7","flash_sale_end","Timestamp","NULLABLE","Hạn Flash Sale"]
                ]
                while len(table.rows) > 1:
                    tbl = table._tbl
                    tr = table.rows[1]._tr
                    tbl.remove(tr)
                for r_data in table_data:
                    row = table.add_row()
                    for i, val in enumerate(r_data):
                        if i < len(row.cells):
                            row.cells[i].text = val
                            
            # Map Users
            elif "User_id" in first_row_text or "Tên tài khoản" in first_row_text:
                table_data = [
                    ["1","id","BigInt","PRIMARY KEY","Mã người dùng"],
                    ["2","name","Varchar(255)","NOT NULL","Tên đầy đủ"],
                    ["3","email","Varchar(255)","NOT NULL","Địa chỉ email"],
                    ["4","password","Varchar(255)","NOT NULL","Mật khẩu"],
                    ["5","role","Varchar(50)","NOT NULL","Quyền (admin/customer)"]
                ]
                while len(table.rows) > 1:
                    tbl = table._tbl
                    tr = table.rows[1]._tr
                    tbl.remove(tr)
                for r_data in table_data:
                    row = table.add_row()
                    for i, val in enumerate(r_data):
                        if i < len(row.cells):
                            row.cells[i].text = val

            # Map Orders
            elif "Order_id" in first_row_text or "Mã đặt hàng" in first_row_text:
                table_data = [
                    ["1","id","BigInt","PRIMARY KEY","Mã đơn hàng"],
                    ["2","user_id","BigInt","FOREIGN KEY","Mã khách hàng"],
                    ["3","total_amount","Decimal(10,2)","NOT NULL","Tổng tiền"],
                    ["4","status","Varchar(50)","NOT NULL","Trạng thái đơn"],
                    ["5","customer_name","Varchar(255)","NULLABLE","Tên người nhận"],
                    ["6","customer_address","Text","NULLABLE","Địa chỉ nhận"],
                    ["7","customer_phone","Varchar(50)","NULLABLE","Số điện thoại"],
                    ["8","voucher_id","BigInt","FOREIGN KEY","Mã Voucher áp dụng"]
                ]
                while len(table.rows) > 1:
                    tbl = table._tbl
                    tr = table.rows[1]._tr
                    tbl.remove(tr)
                for r_data in table_data:
                    row = table.add_row()
                    for i, val in enumerate(r_data):
                        if i < len(row.cells):
                            row.cells[i].text = val

    try:
        doc.save(target_path)
        print("SUCCESS:", target_path)
    except Exception as e:
        print("SAVE FAILED:", e)

if __name__ == "__main__":
    main()
