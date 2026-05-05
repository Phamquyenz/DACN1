import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import urllib.request
import zlib
import base64
import os

b64 = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"

def encode3bytes(b1, b2, b3):
    c1 = b1 >> 2
    c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
    c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
    c4 = b3 & 0x3F
    return b64[c1] + b64[c2] + b64[c3] + b64[c4]

def encode_plantuml(text):
    zlibbed_str = zlib.compress(text.encode('utf-8'))
    compressed_string = zlibbed_str[2:-4]
    res = ""
    i = 0
    length = len(compressed_string)
    while i < length:
        b1 = compressed_string[i] if i < length else 0
        i += 1
        b2 = compressed_string[i] if i < length else 0
        i += 1
        b3 = compressed_string[i] if i < length else 0
        i += 1
        res += encode3bytes(b1, b2, b3)
    return res

def download_diagram(puml_text, filename):
    print("Downloading", filename)
    url = "http://www.plantuml.com/plantuml/png/" + encode_plantuml(puml_text)
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=15) as response, open(filename, 'wb') as out_file:
            out_file.write(response.read())
    except Exception as e:
        print("Failed download:", e)

def create_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

def main():
    doc = docx.Document()
    
    doc.add_paragraph('\n\n\n')
    p = doc.add_paragraph('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ\nKHOA CÔNG NGHỆ THÔNG TIN')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    p_title1 = doc.add_paragraph('BÀI TẬP LỚN\nHỌC PHẦN: ĐỒ ÁN CHUYÊN NGÀNH CÔNG NGHỆ PHẦN MỀM')
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    p_title2 = doc.add_paragraph('TÊN ĐỀ TÀI: XÂY DỰNG WEBSITE QUẢN LÝ BÁN VĂN PHÒNG PHẨM "THE STATIONERY MUSE"')
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_title2.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph('\n\n\n\n\n\n\n\n\n')
    p_title3 = doc.add_paragraph('Sinh viên thực hiện: Phạm Đức Quyền\nGiảng viên hướng dẫn: Mai Văn Linh')
    p_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # MỤC LỤC
    doc.add_heading('MỤC LỤC', 1)
    toc = [
        "CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI",
        "CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG (LARAVEL)",
        "CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ",
        "CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU & ERD",
        "CHƯƠNG V: KẾT LUẬN"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading('CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI', 1)
    doc.add_heading('1.1. Lý do chọn đề tài', 2)
    doc.add_paragraph('Trong thời đại số hóa, thương mại điện tử giúp doanh nghiệp kinh doanh văn phòng phẩm quản lý tồn kho (Stock), đơn hàng, và khuyến mãi hiệu quả hơn so với truyền thống. "The Stationery Muse" được xây dựng để cung cấp công cụ tự động hóa toàn diện từ mua sắm online đến các chương trình chiết khấu ưu đãi khách hàng.')
    doc.add_heading('1.2. Mục tiêu', 2)
    doc.add_paragraph('Xây dựng hệ thống bằng Laravel MVC tinh gọn, kết hợp giao diện với Tailwind CSS, tích hợp thanh toán (VNPAY), kiểm soát chính xác cơ chế áp dụng Mã giảm giá (Vouchers).')

    doc.add_heading('CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG', 1)
    doc.add_paragraph('Dự án dựa trên Framework Laravel 10 (được viết bằng PHP). Khung kiến trúc MVC đảm bảo duy trì logic phần mềm (Controllers) và các mô hình CSDL MySQL rõ ràng. Giao diện tùy ưu nhanh chóng chuyên nghiệp thông qua Tailwind CSS.')

    doc.add_page_break()
    doc.add_heading('CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ', 1)
    
    # Sơ đồ Usecase
    puml_usecase = """@startuml
left to right direction
actor "Khách hàng" as KH
actor "Admin" as AD
package "The Stationery Muse" {
  usecase "Đăng nhập" as UC1
  usecase "Giỏ hàng" as UC2
  usecase "Thanh toán (VNPAY)" as UC3
  usecase "Quản lý / Phân quyền" as UC8
  usecase "Quản lý Sản / Kho" as UC9
  usecase "Quản lý Vouchers" as UC10
  usecase "Xử lý Đơn hàng" as UC11
}
KH --> UC1
KH --> UC2
KH --> UC3
AD --> UC1
AD --> UC8
AD --> UC9
AD --> UC10
AD --> UC11
@enduml"""
    download_diagram(puml_usecase, "uc_diag.png")
    
    doc.add_heading('3.1 Sơ đồ Usecase Tổng quát', 2)
    if os.path.exists("uc_diag.png"):
        doc.add_picture("uc_diag.png", width=Inches(6.0))
        p = doc.add_paragraph('Hình 1: Biểu đồ Usecase hệ thống bán văn phòng phẩm')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Lỗi tải Sơ đồ Use case từ PlantUML]")

    # Sơ đồ Activity
    puml_activity = """@startuml
start
:Thêm Giỏ hàng & Checkout;
:Áp dụng Mã giảm giá (Voucher);
if (Voucher lệ?) then (Có)
  :Trừ theo Giá trị Voucher;
else (Không)
  :Thông báo Lỗi;
endif
:Trừ số lượng Tồn kho của Sản phẩm;
:Tạo Đơn Hàng vào MySQL;
:Mở cổng thanh toán VNPAY;
if (Thanh toán VNPAY?) then (Thành công)
  :Cập nhật trạng thái "Processing";
else (Huỷ / Thất bại)
  :Giao dịch huỷ, hoàn tồn kho;
endif
stop
@enduml"""
    download_diagram(puml_activity, "act_diag.png")

    doc.add_heading('3.2 Sơ đồ Hoạt động (Quy trình Checkout & VNPAY)', 2)
    if os.path.exists("act_diag.png"):
        doc.add_picture("act_diag.png", width=Inches(4.5))
        p = doc.add_paragraph('Hình 2: Sơ đồ hoạt động xử lý vòng lặp đơn hàng')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sơ đồ Sequence
    puml_sequence = """@startuml
actor KH as "Khách hàng"
boundary UI as "Blade View"
control CC as "OrderController"
entity DB as "MySQL"

KH -> UI: Chọn thanh toán & Voucher
UI -> CC: POST Đơn hàng
CC -> DB: Truy vấn chi tiết Voucher
DB --> CC: Thông số giới hạn/hạn mức
alt Áp dụng thành công
    CC -> DB: Update số lượng Sản phẩm (Stock)
    CC -> UI: Khởi tạo Redirect URL (VNPAY)
    UI --> KH: Cổng thanh toán VNPAY
else Hết lượt / Hết thời gian
    CC --> UI: Exception / Flash Error
    UI --> KH: Trả thông báo Lỗi trên View
end
@enduml"""
    download_diagram(puml_sequence, "seq_diag.png")

    doc.add_heading('3.3 Sơ đồ Tuần tự (Xác thực Voucher)', 2)
    if os.path.exists("seq_diag.png"):
        doc.add_picture("seq_diag.png", width=Inches(6.0))
        p = doc.add_paragraph('Hình 3: Sơ đồ tuần tự tương tác Khách hàng với Voucher')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ERD 
    doc.add_heading('CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU & ERD', 1)
    puml_erd = """@startuml
entity "users" as U {
  * id : bigint
  --
  name : varchar
  role : varchar
}
entity "products" as P {
  * id : bigint
  --
  category_id : bigint
  name : varchar
  stock : int
  sale_price : decimal
  flash_sale_end : timestamp
}
entity "vouchers" as V {
  * id : bigint
  --
  code : varchar
  discount_amount : decimal
}
entity "orders" as O {
  * id : bigint
  --
  user_id : bigint
  voucher_id : bigint
  total_amount : decimal
  status : varchar
}
entity "order_items" as OI {
  * id : bigint
  --
  order_id : bigint
  product_id : bigint
  quantity : int
  price : decimal
}
U "1" -- "0..*" O : TẠO >
P "1" -- "0..*" OI : BAO GỒM >
O "1" -- "1..*" OI : LƯU TRỮ >
V "1" -- "0..*" O : ÁP DỤNG MÃ >
@enduml"""
    download_diagram(puml_erd, "erd_diag.png")

    doc.add_heading('4.1 Sơ đồ Thực thể Liên kết (ERD)', 2)
    if os.path.exists("erd_diag.png"):
        doc.add_picture("erd_diag.png", width=Inches(5.0))
        p = doc.add_paragraph('Hình 4: Sơ đồ ERD Hệ thống The Stationery Muse')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.2 Thiết kế Migrations Bảng biểu', 2)
    doc.add_paragraph('Bảng 1: Bảng products (' + 'Sản phẩm'.encode('utf8').decode('utf8') + ')')
    create_table(doc, ['Trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt', 'Khóa chính'],
        ['name', 'Varchar', 'Tên sản phẩm'],
        ['sale_price', 'Decimal', 'Khuyến mãi'],
        ['stock', 'Int', 'Lượng tồn kho thực tế']
    ])

    doc.add_paragraph('\nBảng 2: Bảng vouchers (' + 'Mã giảm giá'.encode('utf8').decode('utf8') + ')')
    create_table(doc, ['Trường', 'Kiểu', 'Ý nghĩa'], [
        ['code', 'Varchar', 'Mã hiển thị (Ví dụ: SALE50)'],
        ['discount_amount', 'Decimal', 'Mức giảm']
    ])

    doc.add_paragraph('\nBảng 3: Bảng order_items (' + 'Chi tiết'.encode('utf8').decode('utf8') + ')')
    create_table(doc, ['Trường', 'Kiểu', 'Ý nghĩa'], [
        ['order_id', 'BigInt', 'Liên kết hệ thống Đơn hàng'],
        ['product_id', 'BigInt', 'Liên kết hệ thống Sản phẩm'],
        ['quantity', 'Int', 'Số lượng xuất kho']
    ])

    doc.add_page_break()
    doc.add_heading('CHƯƠNG V: KẾT LUẬN', 1)
    doc.add_paragraph('Hệ thống đạt mục tiêu thiết lập Website bán hàng The Stationery Muse dựa trên kỹ thuật lập trình MVC bằng Framework Laravel hiện đại. Hệ thống giải quyết tốt bài toán trừ kho sản phẩm và ứng dụng mã giảm giá (Vouchers) tự động.')

    doc.save(r'C:\xampp\htdocs\The Stationery Muse\Báo_Cáo_Siêu_Hoàn_Chỉnh_Có_Hình.docx')

    # Dọn tập tin cũ
    for f in ["uc_diag.png", "act_diag.png", "seq_diag.png", "erd_diag.png"]:
        if os.path.exists(f): os.remove(f)

if __name__ == '__main__':
    main()
