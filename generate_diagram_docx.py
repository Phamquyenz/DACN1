import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import urllib.request
import base64
import zlib
import os

def encode_kroki(text):
    return base64.urlsafe_b64encode(zlib.compress(text.encode('utf-8'), 9)).decode('ascii')

def download_diagram(puml_text, filename):
    url = "https://kroki.io/plantuml/png/" + encode_kroki(puml_text)
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
        out_file.write(response.read())

def create_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

def main():
    doc = docx.Document()

    # Cover Page
    doc.add_paragraph('\n\n\n')
    p = doc.add_paragraph('TRƯỜNG ĐẠI HỌC ...\nKHOA CÔNG NGHỆ THÔNG TIN')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    p_title1 = doc.add_paragraph('BÀI TẬP LỚN\nHỌC PHẦN: ĐỒ ÁN CHUYÊN NGÀNH CÔNG NGHỆ PHẦN MỀM')
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    p_title2 = doc.add_paragraph('TÊN ĐỀ TÀI: XÂY DỰNG WEBSITE QUẢN LÝ BÁN VĂN PHÒNG PHẨM "THE STATIONERY MUSE"')
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_title2.runs:
        run.bold = True
        run.font.size = Pt(16)
    doc.add_paragraph('\n\n\n\n\n\n\n\n\n')
    p_title3 = doc.add_paragraph('Sinh viên thực hiện: ...\nGiảng viên hướng dẫn: ...')
    p_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # MỤC LỤC
    doc.add_heading('MỤC LỤC', 1)
    toc = [
        "CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI",
        "  1.1. Lý do chọn đề tài",
        "  1.2. Mục tiêu của đề tài",
        "CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG (LARAVEL)",
        "CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ",
        "  3.1. Sơ đồ Usecase",
        "  3.2. Sơ đồ Hoạt động",
        "  3.3. Sơ đồ Tuần tự",
        "CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU & ERD",
        "CHƯƠNG V: KẾT LUẬN"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    doc.add_heading('CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI', 1)
    doc.add_heading('1.1. Lý do chọn đề tài', 2)
    doc.add_paragraph('Trong thời đại số hóa, thương mại điện tử giúp doanh nghiệp kinh doanh văn phòng phẩm quản lý tồn kho (Stock), đơn hàng, và khuyến mãi hiệu quả hơn so với truyền thống. "The Stationery Muse" được xây dựng để cung cấp công cụ tự động hóa toàn diện từ mua hàng, sử dụng Voucher, Flash sale, thanh toán online (VNPAY) cho Khách hàng.')
    doc.add_heading('1.2. Mục tiêu', 2)
    doc.add_paragraph('Xây dựng hệ thống bằng Laravel MVC mạnh mẽ, MySQL, và Tailwind CSS với đầy đủ các nghiệp vụ thực tế như xử lí Oversell trong giỏ hàng, giảm giá sản phẩm sâu qua tính năng Voucher và Flash Sale.')

    doc.add_heading('CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG', 1)
    doc.add_paragraph('Dự án dựa trên Framework Laravel (được viết bằng ngôn ngữ PHP). Laravel hỗ trợ ORM (Eloquent) mạnh mẽ giúp xử lý dữ liệu dễ dàng, đi kèm với Blade View Engine làm hệ thống hiển thị dữ liệu web mượt mà. Phía frontend dùng cấu trúc Tailwind CSS để thiết kế diện mạo chuẩn UI/UX và responsive cùng Alpine.js xử lý các component linh hoạt.')

    doc.add_heading('CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ (CÁC SƠ ĐỒ)', 1)
    
    # Sơ đồ Usecase
    puml_usecase = """@startuml
left to right direction
actor "Khách hàng" as KH
actor "Admin" as AD
package "The Stationery Muse" {
  usecase "Đăng ký / Đăng nhập" as UC1
  usecase "Quản lý Giỏ hàng" as UC2
  usecase "Thanh toán (VNPAY & Voucher)" as UC3
  usecase "Quản lý Người dùng" as UC8
  usecase "Quản lý Sản phẩm & Kho" as UC9
  usecase "Quản lý Vouchers" as UC10
  usecase "Xử lý Đơn hàng" as UC11
}
KH --> UC1
KH --> UC2
KH --> UC3
AD --> UC1
AD --> UC8
AD --> UC9
AD --> UC10
AD --> UC11
@enduml"""
    download_diagram(puml_usecase, "uc_diag.png")
    doc.add_heading('3.1 Sơ đồ Usecase Tổng quát', 2)
    doc.add_picture("uc_diag.png", width=Inches(6.0))
    p = doc.add_paragraph('Hình 1: Biểu đồ Usecase hệ thống bán văn phòng phẩm')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sơ đồ Activity
    puml_activity = """@startuml
start
:Khách hàng Checkout đơn hàng;
:Áp dụng Mã giảm giá (Voucher);
if (Voucher hợp lệ?) then (Có)
  :Trừ tiền theo % hoặc trực tiếp;
else (Không)
  :Báo lỗi;
endif
:Lưu Đơn hàng vào CSDL;
:Trừ số lượng Tồn kho (Stock) của Sản phẩm;
:Chuyển hướng cổng thanh toán VNPAY;
if (Khách thanh toán thành công?) then (Có)
  :Cập nhật trạng thái "Processing";
else (Không)
  :Giao dịch huỷ, hoàn lại hàng vào dòng Tồn kho;
endif
stop
@enduml"""
    download_diagram(puml_activity, "act_diag.png")
    doc.add_heading('3.2 Sơ đồ Hoạt động (Quy trình Thanh toán & Đặt hàng)', 2)
    doc.add_picture("act_diag.png", width=Inches(5.0))
    p = doc.add_paragraph('Hình 2: Sơ đồ hoạt động đặt hàng và thanh toán')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sơ đồ Sequence
    puml_sequence = """@startuml
actor KH as "Khách hàng"
boundary UI as "Website (Blade)"
control CC as "CheckoutController"
entity DB as "Database (MySQL)"

KH -> UI: Nhập mã Voucher & Nhấn "Thanh toán"
UI -> CC: Gửi Request (POST)
CC -> DB: Truy vấn VoucherModel
DB --> CC: Thông tin mức giảm giá
alt Áp dụng thành công
    CC -> DB: Lưu Order & OrderItems
    DB --> CC: ID Đơn hàng
    CC -> DB: Trừ đi số lượng Tồn kho (Products.stock)
    CC --> UI: Chuyển khoản VNPAY
    UI --> KH: Cổng thanh toán VNPAY hiện ra
else Voucher hết hạn
    CC --> UI: Cảnh báo Lỗi
    UI --> KH: Hiển thị thông báo thất bại
end
@enduml"""
    download_diagram(puml_sequence, "seq_diag.png")
    doc.add_heading('3.3 Sơ đồ Tuần tự (Quy trình áp dụng Voucher & Database)', 2)
    doc.add_picture("seq_diag.png", width=Inches(6.0))
    p = doc.add_paragraph('Hình 3: Sơ đồ tuần tự xử lý Voucher')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ERD 
    doc.add_heading('CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU & ERD', 1)
    puml_erd = """@startuml
entity "users" as U {
  * id : bigint
  --
  name : varchar
  role : varchar
}
entity "products" as P {
  * id : bigint
  --
  name : varchar
  stock : int
  sale_price : decimal
}
entity "vouchers" as V {
  * id : bigint
  --
  code : varchar
  discount_amount : decimal
}
entity "orders" as O {
  * id : bigint
  --
  user_id : bigint
  voucher_id : bigint
  total_amount : decimal
  status : varchar
}
entity "order_items" as OI {
  * id : bigint
  --
  order_id : bigint
  product_id : bigint
  quantity : int
}
U "1" -- "0..*" O : creates >
P "1" -- "0..*" OI : includes >
O "1" -- "1..*" OI : contains >
V "1" -- "0..*" O : applied in >
@enduml"""
    download_diagram(puml_erd, "erd_diag.png")
    doc.add_heading('4.1 Sơ đồ Thực thể Liên kết (ERD)', 2)
    doc.add_picture("erd_diag.png", width=Inches(5.0))
    p = doc.add_paragraph('Hình 4: Sơ đồ ERD Hệ thống The Stationery Muse')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.2 Các Bảng Database (Migrations)', 2)
    doc.add_paragraph('Bảng 1: Bảng products (Sản phẩm văn phòng phẩm)')
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã sản phẩm'],
        ['category_id', 'BigInt (FK)', 'Liên kết bảng Categories'],
        ['name', 'Varchar', 'Tên sản phẩm'],
        ['sale_price', 'Decimal(10,2)', 'Giá khuyến mãi'],
        ['stock', 'Int', 'Số lượng sản phẩm còn trong kho']
    ])

    doc.add_paragraph('\nBảng 2: Bảng vouchers (Khuyến mãi)')
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã chứng từ'],
        ['code', 'Varchar', 'Mã voucher (VD: SALE50)'],
        ['discount_amount', 'Decimal', 'Mức giảm giá'],
        ['usage_limit', 'Int', 'Giới hạn số lần phát hành']
    ])

    doc.add_paragraph('\nBảng 3: Bảng orders & order_items')
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['order.id', 'BigInt (PK)', 'Mã đơn hàng'],
        ['order.total_amount', 'Decimal', 'Tổng thanh toán'],
        ['order.voucher_id', 'BigInt', 'Voucher đã dùng'],
        ['order_item.quantity', 'Int', 'Số lượng đã mua bảo lưu trong CSDL']
    ])

    doc.add_heading('CHƯƠNG V: KẾT LUẬN', 1)
    doc.add_paragraph('Dự án The Stationery Muse đã vượt qua được các yêu cầu kỹ thuật của một hệ thống thương mại điện tử cấp độ phức tạp bằng nền tảng Laravel mạnh mẽ. Điểm đột phá của dự án là thiết kế quy trình thanh toán chính xác, hệ thống giảm giá Voucher và khả năng trừ sản phẩm trực tiếp khi khách thanh toán để loại bỏ lỗi thất thoát Overselling.')

    doc.save(r'C:\xampp\htdocs\The Stationery Muse\Báo_Cáo_Siêu_Hoàn_Chỉnh_The_Stationery_Muse.docx')

    # cleanup images
    for f in ["uc_diag.png", "act_diag.png", "seq_diag.png", "erd_diag.png"]:
        if os.path.exists(f): os.remove(f)

if __name__ == '__main__':
    main()
