import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    # Customize if needed
    return heading

def create_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

def main():
    doc = docx.Document()

    # Cover Page
    doc.add_paragraph('\n\n\n')
    p = doc.add_paragraph('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ\nKHOA CÔNG NGHỆ THÔNG TIN')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs: pass # run.bold = True
    
    doc.add_paragraph('\n\n')
    p_title1 = doc.add_paragraph('BÀI TẬP LỚN\nHỌC PHẦN: ĐỒ ÁN CHUYÊN NGÀNH CÔNG NGHỆ PHẦN MỀM')
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p_title2 = doc.add_paragraph('TÊN ĐỀ TÀI: XÂY DỰNG WEBSITE BÁN VĂN PHÒNG PHẨM "THE STATIONERY MUSE"')
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_title2.runs:
        run.bold = True
        run.font.size = Pt(16)
        
    doc.add_paragraph('\n\n\n\n\n\n')
    doc.add_page_break()

    # MỤC LỤC
    doc.add_heading('MỤC LỤC', 1)
    toc = [
        "CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI",
        "  1.1. Lý do chọn đề tài",
        "  1.2. Mục tiêu của đề tài",
        "CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG",
        "  2.1. Framework Laravel (PHP)",
        "  2.2. Cơ sở dữ liệu MySQL",
        "  2.3. Frontend (Tailwind CSS, Alpine.js)",
        "CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ",
        "  3.1. Đối tượng người dùng",
        "  3.2. Yêu cầu chức năng",
        "  3.3. Mô hình MVC",
        "CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU",
        "  4.1. Sơ đồ các bảng dữ liệu",
        "CHƯƠNG V: KẾT LUẬN",
        "  5.1. Kết quả đạt được",
        "  5.2. Hướng phát triển"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # CHƯƠNG 1
    doc.add_heading('CHƯƠNG I: TỔNG QUAN VỀ ĐỀ TÀI', 1)
    
    doc.add_heading('1.1. Lý do chọn đề tài', 2)
    doc.add_paragraph('Trong xu thế chuyển đổi số mạnh mẽ, thương mại điện tử đã trở thành một nền tảng không thể thiếu đối với các doanh nghiệp bán lẻ. Văn phòng phẩm là một mặt hàng đặc thù với sự đa dạng mẫu mã, số lượng lớn và nhu cầu cập nhật thường xuyên. Việc quản lý thủ công không còn đáp ứng được yêu cầu về độ chính xác và kiểm soát hàng tồn kho (Stock). Do đó, dự án "The Stationery Muse" được lên ý tưởng và phát triển nhằm số hóa quy trình kinh doanh, xây dựng kênh bán hàng trực tuyến tích hợp thanh toán linh hoạt, đồng thời cung cấp công cụ quản trị mạnh mẽ để quản lý sản phẩm, đơn hàng, và khuyến mãi một cách tự động.')
    
    doc.add_heading('1.2. Mục tiêu của đề tài', 2)
    doc.add_paragraph('- Xây dựng một website bán văn phòng phẩm trực tuyến với giao diện thân thiện, hiện đại.')
    doc.add_paragraph('- Ứng dụng Framework Laravel (PHP) để xây dựng hệ thống bền vững theo chuẩn kiến trúc MVC.')
    doc.add_paragraph('- Thiết kế cơ sở dữ liệu MySQL tối ưu hóa cho các thao tác truy xuất dữ liệu sản phẩm, giỏ hàng (Cart), danh sách yêu thích (Wishlist).')
    doc.add_paragraph('- Tích hợp các tính năng phục vụ bán hàng như: Flash Sale, Hệ thống áp dụng mã giảm giá (Voucher), Thanh toán giao dịch trực tuyến.')

    # CHƯƠNG 2
    doc.add_heading('CHƯƠNG II: CÔNG NGHỆ VÀ NGÔN NGỮ SỬ DỤNG', 1)
    
    doc.add_heading('2.1. Framework Laravel (PHP)', 2)
    doc.add_paragraph('Hệ thống được phát triển hoàn toàn dựa trên Laravel - một Web Framework mã nguồn mở bằng PHP. Laravel cung cấp các công cụ mạnh mẽ như Eloquent ORM cho truy xuất cơ sở dữ liệu, Blade Template cho giao diện, và Artisan Console để xử lý logic Migration. Điều này giúp hệ thống "The Stationery Muse" có được sự phân minh bạch trong mã nguồn (Routing, Middleware, Controllers, Models).')

    doc.add_heading('2.2. Cơ sở dữ liệu MySQL', 2)
    doc.add_paragraph('Hệ thống sử dụng MySQL làm hệ quản trị cơ sở dữ liệu chính. Dữ liệu được cấu trúc chặt chẽ với các ràng buộc khóa ngoại (Foreign Keys) giữa bảng Users, Orders, Products và Categories, đảm bảo tính toàn vẹn dư liệu khi xử lý đơn hàng phức tạp.')

    doc.add_heading('2.3. Frontend (Tailwind CSS, Alpine.js)', 2)
    doc.add_paragraph('Phần giao diện phía người dùng được thiết kế bằng Tailwind CSS – một framework CSS theo dạng utility-first, giúp xây dựng giao diện tùy chỉnh nhanh chóng, đảm bảo tính Responsive (tương thích trên cả PC lẫn Mobile). Kết hợp cùng Alpine.js để xử lý các JavaScript tương tác nhẹ nhàng như đóng/mở menu, popup thông báo, và hiệu ứng giỏ hàng.')

    # CHƯƠNG 3
    doc.add_heading('CHƯƠNG III: YÊU CẦU HỆ THỐNG VÀ THIẾT KẾ', 1)
    
    doc.add_heading('3.1. Đối tượng người dùng', 2)
    doc.add_paragraph('Gồm 2 nhóm tác nhân (Actors) chính: Khách hàng (Customer) và Quản trị viên (Admin).')
    
    doc.add_heading('3.2. Yêu cầu chức năng', 2)
    doc.add_paragraph('Bảng Use Cases chính trị website:')
    create_table(doc, ['Nhóm Người dùng', 'Chức năng (Use case)'], [
        ['Khách hàng', 'Đăng ký, Đăng nhập, Xem/Sửa Profile'],
        ['Khách hàng', 'Tìm kiếm văn phòng phẩm, Lọc theo Danh mục'],
        ['Khách hàng', 'Thêm vào Giỏ hàng (Cart), Danh sách yêu thích (Wishlist)'],
        ['Khách hàng', 'Áp dụng Mã giảm giá (Voucher)'],
        ['Khách hàng', 'Đặt hàng (Checkout) & Thanh toán trực tuyến (VNPAY)'],
        ['Khách hàng', 'Xem bài viết Blog, Đánh giá (Review) sản phẩm'],
        ['Admin', 'Quản lý Tài khoản, Phân quyền người dùng'],
        ['Admin', 'Quản lý Danh mục (Categories), Sản phẩm (Products)'],
        ['Admin', 'Quản lý Đơn hàng (Orders): Duyệt, Hủy, Trạng thái giao hàng'],
        ['Admin', 'Quản lý Mã giảm giá (Vouchers), Chiến dịch Flash Sale'],
        ['Admin', 'Nhập kho hàng (Product Imports)'],
        ['Admin', 'Quản lý Bài viết chia sẻ (Articles)']
    ])

    doc.add_heading('3.3. Mô hình MVC trong hệ thống', 2)
    doc.add_paragraph('Hệ thống hoạt động theo trình tự MVC chặt chẽ:\n- **Model**: Đại diện cho các bảng dữ liệu thực tế: User, Product, Category, Voucher, Cart, Order...\n- **View**: Tập tin .blade.php hiển thị giao diện cho Admin và Khách hàng.\n- **Controller**: ProductController, OrderController, CartController... nhận Request từ người dùng, gọi Model xử lý, và trả về View tương ứng.')


    # CHƯƠNG 4
    doc.add_heading('CHƯƠNG IV: CẤU TRÚC CƠ SỞ DỮ LIỆU', 1)
    doc.add_paragraph('Dưới đây là một số cấu trúc bảng cơ sở dữ liệu cốt lõi đã được trích xuất trực tiếp từ các file Migration trong mã nguồn của hệ thống The Stationery Muse.')

    doc.add_heading('Bảng 1: users (Quản lý Khách hàng và Admin)', 3)
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã người dùng (Tự động tăng)'],
        ['name', 'Varchar', 'Tên đầy đủ'],
        ['email', 'Varchar', 'Email đăng nhập'],
        ['password', 'Varchar', 'Mật khẩu đã mã hoá bcrypt'],
        ['role', 'Enum/Varchar', 'Quyền (admin, customer)']
    ])

    doc.add_paragraph('')
    doc.add_heading('Bảng 2: categories (Quản lý Danh mục sản phẩm)', 3)
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã danh mục'],
        ['name', 'Varchar', 'Tên danh mục ( VD: Bút, Sổ tay)'],
        ['slug', 'Varchar', 'Đường dẫn tĩnh phục vụ SEO']
    ])

    doc.add_paragraph('')
    doc.add_heading('Bảng 3: products (Thông tin Văn phòng phẩm)', 3)
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã sản phẩm'],
        ['category_id', 'BigInt (FK)', 'Liên kết bảng Categories'],
        ['name', 'Varchar', 'Tên sản phẩm'],
        ['description', 'Text', 'Mô tả chi tiết năng/công dụng'],
        ['price', 'Decimal(10,2)', 'Giá bán tiêu chuẩn'],
        ['sale_price', 'Decimal(10,2)', 'Giá khuyến mãi'],
        ['stock', 'Int', 'Số lượng sản phẩm còn trong kho'],
        ['image_url', 'Varchar', 'Hình ảnh sản phẩm'],
        ['flash_sale_end', 'Timestamp', 'Hạn kết thúc nếu có Flash Sale']
    ])

    doc.add_paragraph('')
    doc.add_heading('Bảng 4: vouchers (Quản lý Khuyến mãi)', 3)
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã chứng từ'],
        ['code', 'Varchar', 'Mã người dùng nhập (VD: SALE50)'],
        ['type', 'Varchar', 'Loại (Phần trăm hoặc Trừ tiền trực tiếp)'],
        ['discount_amount', 'Decimal', 'Mức giảm giá'],
        ['usage_limit', 'Int', 'Giới hạn số lần phát hành'],
        ['starts_at', 'Timestamp', 'Ngày bắt đầu chiến dịch'],
        ['expires_at', 'Timestamp', 'Ngày kết thúc chiến dịch']
    ])

    doc.add_paragraph('')
    doc.add_heading('Bảng 5: orders (Quản lý Đơn hàng)', 3)
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã đơn hàng'],
        ['user_id', 'BigInt (FK)', 'Liên kết Người dùng (nếu có đăng nhập)'],
        ['total_amount', 'Decimal(10,2)', 'Tổng giá trị hóa đơn phải trả'],
        ['status', 'Varchar', 'Trạng thái: Pending, Processing, Delivered'],
        ['customer_name', 'Varchar', 'Tên người nhận (Khách lẻ)'],
        ['customer_address','Varchar', 'Địa chỉ giao hàng'],
        ['customer_phone', 'Varchar', 'Điện thoại liên hệ'],
        ['voucher_id', 'BigInt (FK)', 'Mã Voucher đã sử dụng (nếu có)']
    ])

    doc.add_paragraph('')
    doc.add_heading('Bảng 6: order_items (Chi tiết sản phẩm đã bán)', 3)
    create_table(doc, ['Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'], [
        ['id', 'BigInt (PK)', 'Mã ID'],
        ['order_id', 'BigInt (FK)', 'Liên kết bảng Orders'],
        ['product_id', 'BigInt (FK)', 'Liên kết sản phẩm đã mua'],
        ['quantity', 'Int', 'Số lượng hiển thị trong giỏ hàng'],
        ['price', 'Decimal(10,2)', 'Giá bán tại thời điểm tạo đơn']
    ])

    doc.add_page_break()
    # CHƯƠNG 5
    doc.add_heading('CHƯƠNG V: KẾT LUẬN', 1)
    
    doc.add_heading('5.1. Kết quả đạt được', 2)
    doc.add_paragraph('Hệ thống "The Stationery Muse" đã được xây dựng thành công và đáp ứng đầy đủ nghiệp vụ của một nền tảng bán lẻ văn phòng phẩm hiện đại. Công nghệ Laravel bảo đảm mã nguồn được thiết kế chặt chẽ, tối ưu. Các module khó như quản lý tồn kho (chống oversell), xử lý logic áp dụng Voucher, kiểm soát Flash Sale bằng thời gian thực, và giỏ hàng đều hoạt động trơn tru.')

    doc.add_heading('5.2. Hạn chế và Hướng phát triển', 2)
    doc.add_paragraph('- Hạn chế: Việc tạo báo cáo doanh thu dưới dạng PDF dành cho kế toán chưa được tự động hóa. Chức năng đề xuất sản phẩm liên quan vẫn ở mức cơ bản.')
    doc.add_paragraph('- Hướng phát triển: Trong tương lai, hệ thống sẽ tích hợp API của các đơn vị giao hàng Giao Hàng Tiết Kiệm (GHTK) để cập nhật trạng thái vận đơn tự động, triển khai thanh toán đa dạng hơn qua mã QR và Ví Momo, đồng thời tối ưu hóa tính năng gửi Email tự động chăm sóc khách hàng.')

    # Save
    doc.save(r'C:\xampp\htdocs\The Stationery Muse\Bao_Cao_Laravel_The_Stationery_Muse.docx')

if __name__ == '__main__':
    main()
